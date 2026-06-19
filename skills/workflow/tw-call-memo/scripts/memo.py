#!/usr/bin/env python3
"""
tw-call-memo: Generate a Taiwan PE/VC Call Memo from a transcript or audio file.

Pipeline:
  1. Ingest:  --text / --stdin / --file (.txt/.pdf/.docx) / --audio (Whisper)
  2. Extract: Claude LLM with bundled 20-field schema or inferred from --template
  3. Render:  JSON to stdout; optional .docx fill via --docx OUTPUT

See SKILL.md for full documentation.
"""
from __future__ import annotations

import argparse
import io
import json
import os
import re
import subprocess
import sys
from datetime import date
from pathlib import Path

SKILL_ROOT = Path(__file__).resolve().parent.parent
BUNDLED_TEMPLATE = SKILL_ROOT / "data" / "call_memo_template.docx"

# ── Bundled 20-field schema (verbatim from taiwan-company memo_extractor.py) ──

BUNDLED_FIELDS: list[tuple[str, str, str]] = [
    ("deal_source",        "案件來源",                   "例：自行開發、某人介紹。若未提及請填「自行開發」"),
    ("interviewees",       "受訪人",                     "受訪者姓名與職稱，多人以頓號分隔"),
    ("paid_in_capital",    "實收資本額",                  "NT$ 金額，例：5,000萬"),
    ("address",            "地址",                       "公司登記地址或廠址"),
    ("founding_date",      "設立日期",                   "例：2018年 或 2018/03/01"),
    ("underwriter",        "承銷商",                     "輔導券商名稱，若未提及填空"),
    ("auditor",            "會計師事務所",                "簽證會計師事務所，若未提及填空"),
    ("chairman",           "董事長",                     "董事長姓名"),
    ("general_manager",    "總經理",                     "總經理姓名"),
    ("headcount",          "員工人數",                   "數字，例：120人"),
    ("ipo_timeline",       "公開發行及上市櫃時程/募資規劃", "IPO 目標年份、目前募資輪次等時間性資訊"),
    ("investment_terms",   "增資計畫或投資條件",          "本次募資總額、釋出股比、預計 close 時程"),
    ("business_revenue",   "主要業務、產品營收比重",       "核心業務說明及各產品/服務的營收佔比"),
    ("financials",         "財務狀況",                   "近期營收、淨利、年增率等財務數據"),
    ("management_team",    "經營團隊背景",                "創辦人/CEO、CTO、CFO 的背景與經歷"),
    ("board_shareholding", "董監或主要股東持股情形",       "主要股東名稱與持股比例"),
    ("recent_development", "公司發展近況",                "近期重大里程碑、產品進展、合作案"),
    ("major_customers",    "主要銷貨客戶",                "前幾大客戶名稱與佔比"),
    ("major_suppliers",    "主要進貨廠商",                "主要原物料或零件供應商"),
    ("factory_capacity",   "廠房及產能使用情形",          "廠房地點、產能規模、目前使用率"),
    ("competitors",        "國內外主要競爭對手",          "直接競爭者名稱及差異化分析"),
    ("industry_trends",    "產業發展趨勢",                "產業現況、市場規模、未來展望"),
    ("risk_tracking",      "風險評估及追蹤事項",          "主要風險點與需持續追蹤的議題"),
    ("conclusion",         "評估結論與建議",              "訪談整體評估與後續建議行動"),
]

# Extra label-key aliases used by the bundled DOCX template
BUNDLED_LABEL_ALIASES: dict[str, str] = {
    "公司名稱：": "_company_name",
    "案件來源：": "deal_source",
    "受訪人：": "interviewees",
    "實收資本額：": "paid_in_capital",
    "地址：": "address",
    "設立日期：": "founding_date",
    "承銷商：": "underwriter",
    "會計師：": "auditor",
    "董事長：": "chairman",
    "總經理：": "general_manager",
    "員工人數：": "headcount",
    "公開發行及上市櫃時程/募資規劃：": "ipo_timeline",
    "增資計畫或投資條件：": "investment_terms",
    "主要業務、產品營收比重：": "business_revenue",
    "財務狀況：": "financials",
    "經營團隊背景：": "management_team",
    "董監(或主要股東)持股情形：": "board_shareholding",
    "公司發展近況：": "recent_development",
    "主要銷貨客戶：": "major_customers",
    "主要進貨廠商：": "major_suppliers",
    "廠房及產能使用情形：": "factory_capacity",
    "國內外主要競爭對手：": "competitors",
    "產業發展趨勢：": "industry_trends",
    "風險評估及追蹤事項：": "risk_tracking",
    "評估結論與建議：": "conclusion",
}

AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".webm", ".flac", ".aac", ".wma", ".mp4"}


# ── Ingest ────────────────────────────────────────────────────────────────────

def read_transcript(args) -> tuple[str, str | None]:
    """Return (transcript, raw_audio_transcript_or_None)."""
    if args.text:
        return args.text, None
    if args.stdin:
        return sys.stdin.read(), None
    if args.file:
        return _read_file(Path(args.file)), None
    if args.audio:
        text = _transcribe_audio(Path(args.audio), args.whisper_model)
        return text, text
    raise SystemExit("error: one of --text / --stdin / --file / --audio is required")


def _read_file(path: Path) -> str:
    if not path.exists():
        raise SystemExit(f"error: file not found: {path}")
    ext = path.suffix.lower()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".pdf":
        try:
            import fitz
        except ImportError:
            raise SystemExit("error: PDF support requires PyMuPDF: pip install -r scripts/requirements.txt")
        doc = fitz.open(str(path))
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()
    if ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise SystemExit("error: DOCX support requires python-docx: pip install -r scripts/requirements.txt")
        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    raise SystemExit(f"error: unsupported file extension: {ext} (use .txt / .pdf / .docx)")


def _transcribe_audio(path: Path, model_name: str) -> str:
    if not path.exists():
        raise SystemExit(f"error: audio file not found: {path}")
    if path.suffix.lower() not in AUDIO_EXTS:
        raise SystemExit(f"error: unsupported audio format {path.suffix}; supported: {sorted(AUDIO_EXTS)}")
    try:
        import whisper
    except ImportError:
        raise SystemExit(
            "error: audio mode requires openai-whisper.\n"
            "  pip install -r ~/.claude/skills/tw-call-memo/scripts/requirements-audio.txt\n"
            "  Linux: sudo apt install ffmpeg   |   macOS: brew install ffmpeg"
        )
    print(f"[whisper] loading model '{model_name}' (first run downloads ~models)…", file=sys.stderr)
    model = whisper.load_model(model_name)
    print(f"[whisper] transcribing {path.name}…", file=sys.stderr)
    result = model.transcribe(str(path), language="zh", fp16=False)
    return str(result["text"]).strip()


# ── Template ──────────────────────────────────────────────────────────────────

def resolve_template(template_path: str | None) -> tuple[Path, list[tuple[str, str, str]], dict[str, str], str]:
    """
    Returns (template_path, fields, label_aliases, mode):
      - mode = "bundled" or "custom: <path>"
      - fields = [(key, label, description), ...]
      - label_aliases = {full_label_in_template: field_key}
    """
    if not template_path:
        return BUNDLED_TEMPLATE, BUNDLED_FIELDS, BUNDLED_LABEL_ALIASES, "bundled"

    path = Path(template_path).expanduser().resolve()
    if path == BUNDLED_TEMPLATE.resolve():
        return BUNDLED_TEMPLATE, BUNDLED_FIELDS, BUNDLED_LABEL_ALIASES, "bundled"
    if not path.exists():
        raise SystemExit(f"error: template not found: {path}")

    fields, aliases = infer_fields_from_template(path)
    if not fields:
        raise SystemExit(
            f"error: no labelled cells found in {path}.\n"
            "  Custom template must contain a table where each field cell starts with a "
            "first paragraph ending in '：' (full-width colon)."
        )
    return path, fields, aliases, f"custom: {path}"


def infer_fields_from_template(path: Path) -> tuple[list[tuple[str, str, str]], dict[str, str]]:
    """Walk all tables; cells whose first paragraph ends with '：' become field labels."""
    from docx import Document

    doc = Document(str(path))
    fields: list[tuple[str, str, str]] = []
    aliases: dict[str, str] = {}
    seen: set[str] = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                if not paragraphs:
                    continue
                first_line = paragraphs[0].text.strip()
                if not first_line.endswith("："):
                    continue
                label = first_line[:-1].strip()
                if not label or label in seen:
                    continue
                seen.add(label)
                key = label
                fields.append((key, label, label))
                aliases[first_line] = key
    return fields, aliases


# ── LLM extract ───────────────────────────────────────────────────────────────

EXTRACT_PROMPT_TEMPLATE = """你是一位專業的投資分析師助理。以下是與「{company}」的訪談逐字稿。

請從逐字稿中提取以下欄位的資訊，以 JSON 格式回傳。
- 若某欄位在逐字稿中有提及，請整理成清楚的中文句子或條列。
- 若未提及，請回傳空字串 ""。
- 回傳純 JSON，不要加 markdown code block 或其他說明。

需提取的欄位：
{{
{fields_desc}
}}

逐字稿內容：
---
{transcript}
---

請直接回傳 JSON 物件。"""


def extract_fields(
    company: str,
    transcript: str,
    fields: list[tuple[str, str, str]],
    provider: str,
) -> dict[str, str]:
    fields_desc = "\n".join(
        f'  "{key}": "{label}（{desc}）"' for key, label, desc in fields
    )
    prompt = EXTRACT_PROMPT_TEMPLATE.format(
        company=company,
        fields_desc=fields_desc,
        transcript=transcript[:12000],
    )

    raw = _call_claude(prompt, provider)

    raw = re.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=re.MULTILINE)
    raw = re.sub(r"\n?```$", "", raw.strip(), flags=re.MULTILINE)
    m = re.search(r"\{[\s\S]*\}", raw)
    try:
        data = json.loads(m.group()) if m else {}
    except json.JSONDecodeError:
        data = {}

    return {key: str(data.get(key, "")) for key, _, _ in fields}


def _call_claude(prompt: str, provider: str) -> str:
    if provider == "anthropic":
        return _call_anthropic(prompt)
    if provider == "cli":
        return _call_cli(prompt)
    raise SystemExit(f"error: unknown provider {provider}")


def _call_anthropic(prompt: str) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise SystemExit(
            "error: ANTHROPIC_API_KEY not set.\n"
            "  Set it in your environment, or pass --provider cli to use the local claude CLI."
        )
    try:
        import anthropic
    except ImportError:
        raise SystemExit("error: anthropic SDK not installed: pip install -r scripts/requirements.txt")

    client = anthropic.Anthropic(api_key=api_key)
    model = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")
    print(f"[claude] {model}…", file=sys.stderr)
    response = client.messages.create(
        model=model,
        max_tokens=8096,
        messages=[{"role": "user", "content": prompt}],
    )
    return "\n".join(b.text for b in response.content if getattr(b, "type", None) == "text").strip()


def _call_cli(prompt: str) -> str:
    import shutil
    cli = shutil.which("claude") or shutil.which("claude.exe")
    if not cli:
        raise SystemExit("error: claude CLI not found in PATH; set ANTHROPIC_API_KEY and use --provider anthropic")
    print(f"[claude-cli] {cli}…", file=sys.stderr)
    proc = subprocess.run(
        [cli, "-p", prompt],
        capture_output=True, text=True, timeout=180,
    )
    if proc.returncode != 0:
        raise SystemExit(f"error: claude CLI failed: {proc.stderr[:300]}")
    return proc.stdout.strip()


# ── DOCX render ───────────────────────────────────────────────────────────────

def fill_docx(
    template_path: Path,
    output_path: Path,
    company: str,
    fields_data: dict[str, str],
    label_aliases: dict[str, str],
    interview_date: str,
) -> None:
    from docx import Document

    doc = Document(str(template_path))

    # Header paragraphs: replace 訪談日期 placeholder (handles run-splitting)
    if interview_date:
        for para in doc.paragraphs:
            if "訪談日期：" not in para.text:
                continue
            for placeholder in ("2025/X/X", "X/X"):
                if _replace_paragraph_text(para, placeholder, interview_date):
                    break

    def get_value(field_key: str) -> str:
        if field_key == "_company_name":
            return company
        return fields_data.get(field_key, "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                matched_key: str | None = None
                for label, key in label_aliases.items():
                    if cell_text.startswith(label) or cell_text == label.rstrip("：") + "：":
                        matched_key = key
                        break
                if matched_key:
                    _fill_cell(cell, get_value(matched_key))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _replace_paragraph_text(para, search: str, replace: str) -> bool:
    """Replace `search` with `replace` in a paragraph, even when the match spans multiple runs."""
    runs = para.runs
    full = "".join(r.text for r in runs)
    if search not in full:
        return False
    start = full.index(search)
    end = start + len(search)
    pos = 0
    affected: list[tuple[object, int]] = []
    for r in runs:
        r_start = pos
        r_end = pos + len(r.text)
        if r_end > start and r_start < end:
            affected.append((r, r_start))
        pos = r_end
    if not affected:
        return False
    first_run, fr_start = affected[0]
    last_run, lr_start = affected[-1]
    prefix = first_run.text[: start - fr_start]
    suffix = last_run.text[end - lr_start :]
    first_run.text = prefix + replace + suffix
    for r, _ in affected[1:]:
        r.text = ""
    return True


def _fill_cell(cell, value: str) -> None:
    """Keep first paragraph (label), strip prior values, append new value paragraphs."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    tc = cell._tc
    for p in paragraphs[1:]:
        tc.remove(p._p)

    first_p = cell.paragraphs[0]
    label_end_idx = -1
    for i, run in enumerate(first_p.runs):
        if run.bold or (run.text and run.text.strip().endswith("：")):
            label_end_idx = i

    for run in first_p.runs[label_end_idx + 1:]:
        first_p._p.remove(run._r)

    if not value:
        return

    lines = [l for l in value.split("\n") if l.strip()]
    if not lines:
        return

    r = first_p.add_run(" " + lines[0])
    r.bold = False
    for line in lines[1:]:
        new_p = cell.add_paragraph()
        new_p.add_run(line)


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate a Taiwan PE/VC Call Memo from a transcript or audio file.",
    )
    parser.add_argument("--company", required=True, help="Company name (injected into LLM prompt)")

    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("--text", help="Transcript text directly")
    src.add_argument("--stdin", action="store_true", help="Read transcript from stdin")
    src.add_argument("--file", help="Read transcript from .txt / .pdf / .docx")
    src.add_argument("--audio", help="Audio file (mp3/wav/m4a/...) — requires whisper")

    parser.add_argument("--template", help="Custom DOCX template (default: bundled)")
    parser.add_argument("--docx", help="Write filled DOCX to this path")
    parser.add_argument("--date", default=date.today().strftime("%Y/%m/%d"),
                        help="Interview date YYYY/MM/DD (default: today)")
    parser.add_argument("--provider", choices=["anthropic", "cli"], default="anthropic",
                        help="LLM provider")
    parser.add_argument("--whisper-model", default="small",
                        choices=["tiny", "base", "small", "medium", "large"],
                        help="Whisper model size (default: small)")
    parser.add_argument("--json-only", action="store_true",
                        help="Suppress stderr summary; only emit JSON to stdout")

    args = parser.parse_args()

    transcript, audio_text = read_transcript(args)
    if not transcript.strip():
        raise SystemExit("error: transcript is empty")

    template_path, fields, aliases, mode = resolve_template(args.template)

    if not args.json_only:
        print(f"[ingest] {len(transcript)} chars  template={mode}  fields={len(fields)}",
              file=sys.stderr)

    fields_data = extract_fields(args.company, transcript, fields, args.provider)

    docx_output: str | None = None
    if args.docx:
        out_path = Path(args.docx).expanduser().resolve()
        fill_docx(template_path, out_path, args.company, fields_data, aliases, args.date)
        docx_output = str(out_path)
        if not args.json_only:
            print(f"[docx]  wrote {out_path}", file=sys.stderr)

    result = {
        "company": args.company,
        "interview_date": args.date,
        "template": mode,
        "transcript": audio_text,
        "fields": fields_data,
        "docx_output": docx_output,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))

    if not args.json_only:
        non_empty = sum(1 for v in fields_data.values() if v)
        print(f"[done]  {non_empty}/{len(fields_data)} fields populated", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
